import sys
import speech_recognition as sr
import traceback
from pathlib import Path
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QPushButton, 
                            QVBoxLayout, QWidget, QFileDialog, QTextEdit,
                            QProgressBar, QMessageBox, QHBoxLayout, QLabel,
                            QStatusBar, QFrame, QComboBox, QSpinBox, QMenu,
                            QTableWidget, QTableWidgetItem, QDialog, QLineEdit,
                            QDialogButtonBox, QFormLayout, QTabWidget, QScrollArea)
from PyQt6.QtCore import Qt, QSize, QByteArray, QTimer
from PyQt6.QtGui import QFont, QIcon, QAction, QTextCursor
import os
from converter import AudioConverterThread
from styles import StyleSheet
from history import ConversionHistory
from config import AppConfig

# Comprobar si python-docx está disponible
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    HAS_REPORTLAB = True
except:
    HAS_REPORTLAB = False

class SearchReplaceDialog(QDialog):
    """Diálogo para buscar y reemplazar texto"""
    
    def __init__(self, parent, text_edit):
        super().__init__(parent)
        self.text_edit = text_edit
        self.setWindowTitle('Buscar y reemplazar')
        self.setGeometry(200, 200, 400, 150)
        self.setStyleSheet(StyleSheet.get_styles())
        
        layout = QFormLayout()
        
        self.search_input = QLineEdit()
        self.replace_input = QLineEdit()
        
        layout.addRow('Buscar:', self.search_input)
        layout.addRow('Reemplazar:', self.replace_input)
        
        button_layout = QHBoxLayout()
        
        find_btn = QPushButton('Buscar')
        replace_btn = QPushButton('Reemplazar')
        replace_all_btn = QPushButton('Reemplazar todo')
        close_btn = QPushButton('Cerrar')
        
        find_btn.clicked.connect(self.find_text)
        replace_btn.clicked.connect(self.replace_text)
        replace_all_btn.clicked.connect(self.replace_all)
        close_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(find_btn)
        button_layout.addWidget(replace_btn)
        button_layout.addWidget(replace_all_btn)
        button_layout.addWidget(close_btn)
        
        layout.addRow(button_layout)
        
        self.setLayout(layout)
    
    def find_text(self):
        text = self.search_input.text()
        if text:
            content = self.text_edit.toPlainText()
            if text in content:
                self.text_edit.find(text)
    
    def replace_text(self):
        search_text = self.search_input.text()
        replace_text = self.replace_input.text()
        
        if search_text:
            content = self.text_edit.toPlainText()
            new_content = content.replace(search_text, replace_text, 1)
            self.text_edit.setText(new_content)
    
    def replace_all(self):
        search_text = self.search_input.text()
        replace_text = self.replace_input.text()
        
        if search_text:
            content = self.text_edit.toPlainText()
            count = content.count(search_text)
            new_content = content.replace(search_text, replace_text)
            self.text_edit.setText(new_content)
            QMessageBox.information(self, "Resultado", 
                                  f"Se reemplazaron {count} coincidencias")

class SettingsDialog(QDialog):
    """Diálogo de configuración"""
    
    def __init__(self, parent, config):
        super().__init__(parent)
        self.config = config
        self.setWindowTitle('Configuración')
        self.setGeometry(200, 200, 400, 300)
        self.setStyleSheet(StyleSheet.get_styles())
        
        layout = QFormLayout()
        
        # Selector de idioma
        self.language_combo = QComboBox()
        for lang_code, lang_name in AppConfig.SUPPORTED_LANGUAGES.items():
            self.language_combo.addItem(lang_name, lang_code)
        
        current_lang = config.get_language()
        index = self.language_combo.findData(current_lang)
        if index >= 0:
            self.language_combo.setCurrentIndex(index)
        
        layout.addRow('Idioma:', self.language_combo)
        

        
        # Botones
        button_layout = QHBoxLayout()
        save_btn = QPushButton('Guardar')
        cancel_btn = QPushButton('Cancelar')
        
        save_btn.clicked.connect(self.accept)
        cancel_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(save_btn)
        button_layout.addWidget(cancel_btn)
        
        layout.addRow(button_layout)
        
        self.setLayout(layout)
    
    def get_language(self):
        return self.language_combo.currentData()
    
    def get_duration(self):
        return None  # Sin límite

class HistoryDialog(QDialog):
    """Diálogo para ver el historial de conversiones"""
    
    def __init__(self, parent, history):
        super().__init__(parent)
        self.history = history
        self.selected_text = None
        self.setWindowTitle('Historial de conversiones')
        self.setGeometry(100, 100, 800, 500)
        self.setStyleSheet(StyleSheet.get_styles())
        
        layout = QVBoxLayout()
        
        # Tabla de historial
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['Fecha', 'Archivo', 'Vista previa', 'Palabras'])
        self.table.setColumnWidth(0, 150)
        self.table.setColumnWidth(1, 150)
        self.table.setColumnWidth(2, 350)
        self.table.setColumnWidth(3, 80)
        
        # Llenar la tabla
        conversions = history.get_history()
        for i, conv in enumerate(conversions):
            self.table.insertRow(i)
            
            timestamp = datetime.fromisoformat(conv['timestamp']).strftime('%d/%m/%Y %H:%M')
            self.table.setItem(i, 0, QTableWidgetItem(timestamp))
            self.table.setItem(i, 1, QTableWidgetItem(conv['filename']))
            self.table.setItem(i, 2, QTableWidgetItem(conv['text_preview']))
            
            word_count = len(conv['full_text'].split())
            self.table.setItem(i, 3, QTableWidgetItem(str(word_count)))
        
        layout.addWidget(self.table)
        
        # Botones
        button_layout = QHBoxLayout()
        
        restore_btn = QPushButton('Restaurar')
        clear_btn = QPushButton('Limpiar historial')
        close_btn = QPushButton('Cerrar')
        
        restore_btn.clicked.connect(self.restore_selection)
        clear_btn.clicked.connect(self.clear_history)
        close_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(restore_btn)
        button_layout.addWidget(clear_btn)
        button_layout.addStretch()
        button_layout.addWidget(close_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
    
    def restore_selection(self):
        current_row = self.table.currentRow()
        if current_row >= 0:
            conversion = self.history.get_conversion_by_index(current_row)
            self.selected_text = conversion['full_text']
            self.accept()
    
    def clear_history(self):
        reply = QMessageBox.question(self, 'Confirmar', 
                                    '¿Desea limpiar todo el historial?',
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.history.clear_history()
            self.table.setRowCount(0)

class AudioConverterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.config = AppConfig()
        self.history = ConversionHistory()
        self.setWindowOpacity(0.98)
        self.conversion_data = None
        
        self.initUI()
        self.audio_file = None
        self.converter_thread = None
        self.file_queue = []
        self.is_editing_text = False
        
        try:
            self.recognizer = sr.Recognizer()
            self.status_bar.showMessage('Sistema inicializado correctamente')
        except Exception as e:
            QMessageBox.critical(self, "Error de inicialización", str(e))
            sys.exit(1)
    
    def initUI(self):
        self.setWindowTitle('Convertidor de audio a texto')
        self.setGeometry(100, 100, 1200, 800)
        self.setStyleSheet(StyleSheet.get_styles())
        self.showMaximized()
        self.setWindowIcon(QIcon("favicon.ico"))
        
        # Crear menú
        self.create_menu()

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(20, 20, 20, 20)

        separator = QFrame()
        separator.setObjectName("separator")
        separator.setFrameShape(QFrame.Shape.HLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        main_layout.addWidget(separator)

        # Panel de configuración
        config_container = QWidget()
        config_container.setStyleSheet(StyleSheet.get_button_container_style())
        config_layout = QHBoxLayout(config_container)
        config_layout.setSpacing(10)
        
        config_layout.addWidget(QLabel('Idioma:'))
        self.language_combo = QComboBox()
        for lang_code, lang_name in AppConfig.SUPPORTED_LANGUAGES.items():
            self.language_combo.addItem(lang_name, lang_code)
        current_lang = self.config.get_language()
        index = self.language_combo.findData(current_lang)
        if index >= 0:
            self.language_combo.setCurrentIndex(index)
        config_layout.addWidget(self.language_combo)
        
        config_layout.addStretch()
        
        main_layout.addWidget(config_container)

        # Panel de botones
        button_container = QWidget()
        button_container.setStyleSheet(StyleSheet.get_button_container_style())
        button_layout = QHBoxLayout(button_container)
        button_layout.setSpacing(10)
        
        self.load_button = QPushButton('Cargar audio')
        self.convert_button = QPushButton('Convertir')
        self.copy_button = QPushButton('Copiar')
        self.save_button = QPushButton('Guardar')
        self.cancel_button = QPushButton('Cancelar')
        self.edit_button = QPushButton('Editar texto')

        for button in [self.load_button, self.convert_button, self.copy_button,
                    self.save_button, self.cancel_button, self.edit_button]:
            button.setMinimumSize(QSize(120, 45))
            button.setFont(QFont('Segoe UI', 10))
            button_layout.addWidget(button)

        self.convert_button.setEnabled(False)
        self.copy_button.setEnabled(False)
        self.save_button.setEnabled(False)
        self.cancel_button.setEnabled(False)
        self.edit_button.setEnabled(False)

        self.load_button.clicked.connect(self.load_audio)
        self.convert_button.clicked.connect(self.convert_audio)
        self.copy_button.clicked.connect(self.copy_to_clipboard)
        self.save_button.clicked.connect(self.save_text)
        self.cancel_button.clicked.connect(self.cancel_conversion)
        self.edit_button.clicked.connect(self.toggle_edit_text)

        main_layout.addWidget(button_container)

        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setMinimumHeight(20)
        self.progress_bar.setFont(QFont('Segoe UI', 9))
        main_layout.addWidget(self.progress_bar)

        # Contenedor de información
        info_container = QWidget()
        info_container.setStyleSheet(StyleSheet.get_button_container_style())
        info_layout = QHBoxLayout(info_container)
        
        self.word_count_label = QLabel('Palabras: 0')
        self.char_count_label = QLabel('Caracteres: 0')
        self.duration_label = QLabel('Duración: -')
        self.lang_label = QLabel('Idioma: -')
        self.confidence_label = QLabel('Confianza: -')
        
        for label in [self.word_count_label, self.char_count_label, 
                     self.duration_label, self.lang_label, self.confidence_label]:
            label.setFont(QFont('Segoe UI', 9))
            info_layout.addWidget(label)
        
        main_layout.addWidget(info_container)

        text_container = QWidget()
        text_container.setStyleSheet(StyleSheet.get_text_container_style())
        text_layout = QVBoxLayout(text_container)
        
        self.text_area = QTextEdit()
        self.text_area.setReadOnly(True)
        self.text_area.setMinimumHeight(400)
        self.text_area.setFont(QFont('Segoe UI', 11))
        self.text_area.setAcceptRichText(False)
        
        text_layout.addWidget(self.text_area)
        
        main_layout.addWidget(text_container)

        self.status_bar = QStatusBar()
        self.status_bar.setFont(QFont('Segoe UI', 9))
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage('Listo')
        
        # Permitir drag & drop en la ventana
        self.setAcceptDrops(True)
    
    def create_menu(self):
        """Crea el menú de la aplicación"""
        menubar = self.menuBar()
        
        # Menú Archivo
        file_menu = menubar.addMenu('&Archivo')
        
        open_action = QAction('&Abrir audio\tCtrl+O', self)
        open_action.triggered.connect(self.load_audio)
        file_menu.addAction(open_action)
        
        file_menu.addSeparator()
        
        save_action = QAction('&Guardar texto\tCtrl+S', self)
        save_action.triggered.connect(self.save_text)
        file_menu.addAction(save_action)
        
        save_menu = file_menu.addMenu('Guardar como...')
        
        if HAS_DOCX:
            save_docx = QAction('Documento Word (.docx)', self)
            save_docx.triggered.connect(lambda: self.export_format('docx'))
            save_menu.addAction(save_docx)
        
        if HAS_REPORTLAB:
            save_pdf = QAction('Documento PDF (.pdf)', self)
            save_pdf.triggered.connect(lambda: self.export_format('pdf'))
            save_menu.addAction(save_pdf)
        
        save_md = QAction('Markdown (.md)', self)
        save_md.triggered.connect(lambda: self.export_format('md'))
        save_menu.addAction(save_md)
        
        file_menu.addSeparator()
        
        exit_action = QAction('&Salir\tCtrl+Q', self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Menú Editar
        edit_menu = menubar.addMenu('&Editar')
        
        copy_action = QAction('&Copiar\tCtrl+C', self)
        copy_action.triggered.connect(self.copy_to_clipboard)
        edit_menu.addAction(copy_action)
        
        find_action = QAction('&Buscar y reemplazar\tCtrl+H', self)
        find_action.triggered.connect(self.open_search_replace)
        edit_menu.addAction(find_action)
        
        # Menú Ver
        view_menu = menubar.addMenu('&Ver')
        
        history_action = QAction('&Historial de conversiones', self)
        history_action.triggered.connect(self.show_history)
        view_menu.addAction(history_action)
        
        # Menú Herramientas
        tools_menu = menubar.addMenu('&Herramientas')
        
        settings_action = QAction('&Configuración', self)
        settings_action.triggered.connect(self.open_settings)
        tools_menu.addAction(settings_action)
        
        # Menú Ayuda
        help_menu = menubar.addMenu('Ay&uda')
        
        about_action = QAction('&Acerca de', self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def dragEnterEvent(self, event):
        """Permite arrastrar archivos a la ventana"""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
    
    def dropEvent(self, event):
        """Maneja el evento de soltar archivos"""
        files = [url.toLocalFile() for url in event.mimeData().urls()]
        for file_path in files:
            if file_path.lower().endswith(('.mp3', '.wav', '.m4a')):
                self.audio_file = file_path
                self.convert_button.setEnabled(True)
                self.text_area.setText(f"Archivo cargado: {Path(file_path).name}")
                self.status_bar.showMessage('Archivo cargado correctamente')
                break
    
    def center_window(self):
        """Centra la ventana en la pantalla actual"""
        screen = QApplication.primaryScreen().availableGeometry()
        x = (screen.width() - self.width()) // 2
        y = (screen.height() - self.height()) // 2
        self.move(x, y)
    
    def load_audio(self):
        try:
            file_name, _ = QFileDialog.getOpenFileName(
                self,
                "Seleccionar archivo de audio",
                self.config.get('last_path', str(Path.home())),
                "Audio Files (*.mp3 *.wav *.m4a)"
            )

            if file_name:
                self.audio_file = file_name
                self.config.set('last_path', str(Path(file_name).parent))
                self.convert_button.setEnabled(True)
                self.text_area.setText(f"Archivo cargado: {Path(file_name).name}")
                self.status_bar.showMessage('Archivo cargado correctamente')
        except Exception as e:
            self.show_error("Error al cargar el archivo", str(e))
    
    def convert_audio(self):
        if not self.audio_file:
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.update_ui_state(is_converting=True)

        language = self.language_combo.currentData()
        self.converter_thread = AudioConverterThread(self.audio_file, self.recognizer, language)
        self.converter_thread.progress.connect(self.update_progress)
        self.converter_thread.status.connect(self.update_status)
        self.converter_thread.finished.connect(self.conversion_finished)
        self.converter_thread.error.connect(self.show_error)
        self.converter_thread.start()
    
    def update_ui_state(self, is_converting=False):
        """Actualiza el estado de los botones"""
        self.convert_button.setEnabled(not is_converting and self.audio_file is not None)
        self.load_button.setEnabled(not is_converting)
        self.cancel_button.setEnabled(is_converting)
        has_text = bool(self.text_area.toPlainText().strip())
        self.save_button.setEnabled(not is_converting and has_text)
        self.copy_button.setEnabled(not is_converting and has_text)
        self.edit_button.setEnabled(not is_converting and has_text)
        self.language_combo.setEnabled(not is_converting)
    
    def conversion_finished(self, result):
        """Maneja el fin de la conversión"""
        try:
            if isinstance(result, dict) and 'text' in result:
                text = result['text']
                self.conversion_data = result
                self.text_area.setText(text)
                self.status_bar.showMessage('Conversión completada exitosamente')
                
                # Actualizar información
                self.update_info_labels(result)
                
                # Guardar en historial
                self.history.add_conversion(
                    self.audio_file,
                    text,
                    result.get('duration', 0),
                    result.get('language', 'es-ES'),
                    result.get('confidence', 0)
                )
            else:
                self.text_area.setText("No se pudo extraer texto del audio")
                self.status_bar.showMessage('Conversión completada sin resultados')
            
            self.reset_ui()
            
        except Exception as e:
            self.show_error("Error al finalizar la conversión", str(e))
    
    def update_info_labels(self, result):
        """Actualiza las etiquetas de información"""
        if result:
            duration = result.get('duration', 0)
            word_count = result.get('word_count', 0)
            char_count = len(result.get('text', ''))
            language = result.get('language', '-')
            confidence = result.get('confidence', 0)
            
            self.word_count_label.setText(f'Palabras: {word_count}')
            self.char_count_label.setText(f'Caracteres: {char_count}')
            self.duration_label.setText(f'Duración: {duration:.1f}s' if duration > 0 else 'Duración: -')
            self.lang_label.setText(f'Idioma: {language}')
            self.confidence_label.setText(f'Confianza: {confidence*100:.0f}%' if confidence > 0 else 'Confianza: -')
    
    def reset_ui(self):
        """Resetea la interfaz"""
        try:
            self.progress_bar.setVisible(False)
            self.progress_bar.setValue(0)
            self.convert_button.setEnabled(True)
            self.load_button.setEnabled(True)
            self.cancel_button.setEnabled(False)
            has_text = bool(self.text_area.toPlainText().strip())
            self.save_button.setEnabled(has_text)
            self.copy_button.setEnabled(has_text)
            self.edit_button.setEnabled(has_text)
            self.language_combo.setEnabled(True)
        except Exception as e:
            self.show_error("Error al resetear la interfaz", str(e))
    
    def cancel_conversion(self):
        """Cancela la conversión en curso"""
        if self.converter_thread and self.converter_thread.isRunning():
            self.converter_thread.cancel()
            self.converter_thread.wait()
            self.reset_ui()
            self.text_area.setText("Conversión cancelada")
            self.status_bar.showMessage('Conversión cancelada')
    
    def update_progress(self, value):
        self.progress_bar.setValue(value)
    
    def update_status(self, message):
        self.status_bar.showMessage(message)
    
    def copy_to_clipboard(self):
        """Copia el texto al portapapeles"""
        try:
            text = self.text_area.toPlainText()
            if text:
                clipboard = QApplication.clipboard()
                clipboard.setText(text)
                self.status_bar.showMessage('Texto copiado al portapapeles')
                
                # Mostrar notificación temporal
                QTimer.singleShot(3000, lambda: self.status_bar.showMessage('Listo'))
        except Exception as e:
            self.show_error("Error al copiar", str(e))
    
    def toggle_edit_text(self):
        """Habilita/deshabilita la edición de texto"""
        self.is_editing_text = not self.is_editing_text
        self.text_area.setReadOnly(not self.is_editing_text)
        
        if self.is_editing_text:
            self.edit_button.setText('Finalizar edición')
            self.status_bar.showMessage('Modo edición activado')
        else:
            self.edit_button.setText('Editar texto')
            self.status_bar.showMessage('Modo edición desactivado')
    
    def open_search_replace(self):
        """Abre el diálogo de búsqueda y reemplazo"""
        if self.text_area.toPlainText():
            dialog = SearchReplaceDialog(self, self.text_area)
            dialog.exec()
    
    def save_text(self):
        """Guarda el texto en archivo"""
        try:
            file_name, _ = QFileDialog.getSaveFileName(
                self,
                "Guardar texto",
                str(Path.home() / "texto_convertido.txt"),
                "Archivo de texto (*.txt);;Todos los archivos (*.*)"
            )

            if file_name:
                with open(file_name, 'w', encoding='utf-8') as f:
                    f.write(self.text_area.toPlainText())
                self.status_bar.showMessage('Archivo guardado correctamente')
                QMessageBox.information(self, "Éxito", 
                                      "El texto se ha guardado correctamente")
        except Exception as e:
            self.show_error("Error al guardar el archivo", str(e))
    
    def export_format(self, format_type):
        """Exporta el texto en diferentes formatos"""
        if not self.text_area.toPlainText():
            QMessageBox.warning(self, "Advertencia", "No hay texto para exportar")
            return
        
        try:
            if format_type == 'docx':
                if not HAS_DOCX:
                    QMessageBox.warning(self, "Advertencia", "python-docx no está instalado")
                    return
                
                file_name, _ = QFileDialog.getSaveFileName(
                    self, "Guardar como word", str(Path.home()),
                    "Word Files (*.docx)"
                )
                
                if file_name:
                    doc = Document()
                    doc.add_heading('Transcripción de audio', 0)
                    doc.add_paragraph(self.text_area.toPlainText())
                    doc.save(file_name)
                    self.status_bar.showMessage('Documento word guardado')
                    QMessageBox.information(self, "Éxito", "Archivo word guardado correctamente")
            
            elif format_type == 'pdf':
                if not HAS_REPORTLAB:
                    QMessageBox.warning(self, "Advertencia", "reportlab no está instalado")
                    return
                
                file_name, _ = QFileDialog.getSaveFileName(
                    self, "Guardar como PDF", str(Path.home()),
                    "PDF Files (*.pdf)"
                )
                
                if file_name:
                    doc = SimpleDocTemplate(file_name, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                    story.append(Paragraph("Transcripción de audio", styles['Heading1']))
                    story.append(Spacer(1, 12))
                    story.append(Paragraph(self.text_area.toPlainText().replace('\n', '<br/>'), 
                                         styles['BodyText']))
                    doc.build(story)
                    self.status_bar.showMessage('PDF guardado')
                    QMessageBox.information(self, "Éxito", "PDF guardado correctamente")
            
            elif format_type == 'md':
                file_name, _ = QFileDialog.getSaveFileName(
                    self, "Guardar como markdown", str(Path.home()),
                    "Markdown Files (*.md)"
                )
                
                if file_name:
                    with open(file_name, 'w', encoding='utf-8') as f:
                        f.write('# Transcripción de audio\n\n')
                        f.write(self.text_area.toPlainText())
                    self.status_bar.showMessage('Markdown guardado')
                    QMessageBox.information(self, "Éxito", "Markdown guardado correctamente")
        
        except Exception as e:
            self.show_error(f"Error al exportar como {format_type}", str(e))
    
    def show_history(self):
        """Muestra el diálogo de historial"""
        dialog = HistoryDialog(self, self.history)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            if dialog.selected_text:
                self.text_area.setText(dialog.selected_text)
                self.status_bar.showMessage('Conversión restaurada del historial')
    
    def open_settings(self):
        """Abre el diálogo de configuración"""
        dialog = SettingsDialog(self, self.config)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_language = dialog.get_language()
            new_duration = dialog.get_duration()
            
            self.config.set_language(new_language)
            self.config.set_max_duration(new_duration)
            
            # Actualizar combo de idioma
            index = self.language_combo.findData(new_language)
            if index >= 0:
                self.language_combo.setCurrentIndex(index)
            
            self.status_bar.showMessage('Configuración guardada')
    
    def show_about(self):
        """Muestra el diálogo de información"""
        about_text = """
        <h2>Convertidor de audio a texto</h2>
        <p><b>Versión:</b> 2.0</p>
        <p><b>Descripción:</b> Aplicación para convertir archivos de audio a texto usando reconocimiento de voz.</p>
        
        <p><b>Características:</b></p>
        <ul>
            <li>Soporte para múltiples idiomas</li>
            <li>Exportación a PDF, DOCX y Markdown</li>
            <li>Historial de conversiones</li>
            <li>Búsqueda y reemplazo de texto</li>
            <li>Soporte para Drag & Drop</li>
        </ul>
        
        <p><b>Dependencias:</b></p>
        <ul>
            <li>SpeechRecognition</li>
            <li>moviepy</li>
            <li>PyQt6</li>
        </ul>
        """
        
        QMessageBox.about(self, "Acerca de", about_text)
    
    def show_error(self, title, message=None):
        """Muestra un mensaje de error"""
        if message is None:
            message = title
            title = "Error"

        error_box = QMessageBox(self)
        error_box.setStyleSheet(StyleSheet.get_error_box_style())
        error_box.setIcon(QMessageBox.Icon.Critical)
        error_box.setWindowTitle(title)
        error_box.setText(message)
        error_box.exec()
        self.status_bar.showMessage(f'Error: {title}')
    
    def closeEvent(self, event):
        """Maneja el evento de cierre"""
        if self.converter_thread and self.converter_thread.isRunning():
            reply = QMessageBox.question(
                self, 'Confirmar salida',
                '¿Está seguro de que desea salir? La conversión en curso se cancelará.',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                try:
                    self.converter_thread.cancel()
                    self.converter_thread.wait()
                except Exception as e:
                    print(f"Error al cancelar la conversión: {e}")
                finally:
                    event.accept()
            else:
                event.ignore()
        else:
            event.accept()

def main():
    app = QApplication(sys.argv)
    try:
        app_path = Path(__file__).parent
        os.chdir(str(app_path))
        
        try:
            recognizer = sr.Recognizer()
        except Exception as e:
            QMessageBox.warning(
                None, 
                "Error de inicialización",
                f"Error al inicializar el sistema:\n{str(e)}\n\n"
                "Verifica que todas las dependencias están correctamente instaladas."
            )
            return

        ex = AudioConverterApp()
        ex.show()
        sys.exit(app.exec())
    except Exception as e:
        QMessageBox.critical(
            None, 
            "Error Fatal", 
            f"Error al iniciar la aplicación:\n{str(e)}\n\n"
            f"Detalles técnicos:\n{traceback.format_exc()}"
        )
        sys.exit(1)

if __name__ == '__main__':
    main()
